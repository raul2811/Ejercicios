import os
import tempfile
from pyonedrive import LiveAuth
from pyonedrive import OneDrive
import telegram
from telegram import Update
from telegram.ext import Updater, CommandHandler, MessageHandler, Filters, ConversationHandler, CallbackContext
from docx import Document
import docx2pdf

# URL del archivo Word en OneDrive
onedrive_word_url = "https://upanama-my.sharepoint.com/:f:/g/personal/raul_serrano01_up_ac_pa/ElkPZ0uNSgdAkUpQXstBndsBL0vsRRUhu5ztz6SJ3HCmug?e=gWQhAQ"
onedrive_pdf_url = "https://upanama-my.sharepoint.com/:f:/g/personal/raul_serrano01_up_ac_pa/EiuyFMx-0MBDk-N2W8lNld4BDYhQleqjn_MM1xu4eWe1NA?e=ujFiLG"
onedrive_word_plantilla="https://upanama-my.sharepoint.com/:w:/g/personal/raul_serrano01_up_ac_pa/ERHHLKqQVltOoAUJcdiZoSkBfklB_-0I9RpQ1OnySCvOOw?e=AIfPsW"

# Estados de la conversación
GENDER, NAME, CEDULA, DATE, START_TIME, END_TIME, SIGN_DATE = range(7)

# Diccionario para la conversión de cédula
cedula_conversion = {
    "-": "Guion",
    "0": "Cero",
    "1": "Uno",
    "2": "Dos",
    "3": "Tres",
    "4": "Cuatro",
    "5": "Cinco",
    "6": "Seis",
    "7": "Siete",
    "8": "Ocho",
    "9": "Nueve"
}

# Función para convertir una cédula en letras
def convertir_cedula_a_letras(cedula):
    letras = [cedula_conversion.get(caracter, caracter) for caracter in cedula]
    return " ".join(letras)
# Función para guardar el archivo en OneDrive

def guardar_en_onedrive(archivo_local, url_destino):
    try:
        onedrive = OneDriveAPI.from_config('config.json')  # Asegúrate de tener un archivo de configuración adecuado
        onedrive.auth()
        with open(archivo_local, "rb") as archivo:
            onedrive.upload(archivo, url_destino)
        print(f"Archivo guardado en OneDrive: {url_destino}")
    except Exception as e:
        print(f"Error al guardar en OneDrive: {str(e)}")
        
# Variable global para rastrear si el bot está en funcionamiento
bot_en_funcionamiento = True

# Función para iniciar la conversación
def start(update: Update, context: CallbackContext):
    global bot_en_funcionamiento
    if bot_en_funcionamiento:
        reply_text = "¡Hola! Soy tu bot de contrato. Para detenerme, usa /stop."
    else:
        reply_text = "El bot ya se ha detenido. Para volver a iniciarlo, usa /start."
    update.message.reply_text(reply_text)

    # Solicita al usuario que elija el género
    reply_text = "Hola, soy tu bot de contrato. Por favor, elige el género: (1) Hombre, (2) Mujer"
    update.message.reply_text(reply_text)
    return GENDER

# Función para detener el bot
def stop(update: Update, context: CallbackContext):
    global bot_en_funcionamiento
    bot_en_funcionamiento = False
    update.message.reply_text("El bot se ha detenido. Para volver a iniciarlo, usa /start.")

# Función para manejar la elección de género
def choose_gender(update, context):
    global bot_en_funcionamiento
    if not bot_en_funcionamiento:
        update.message.reply_text("El bot se ha detenido. Para volver a iniciarlo, usa /start.")
        return ConversationHandler.END

    try:
        gender = int(update.message.text)
        context.user_data['gender'] = gender
        if gender == 1:
            context.user_data['saludo'] = "El Modelo"
            context.user_data['seno'] = "el señor"
            context.user_data['gen'] = "hombre"
            context.user_data['pana'] = "panameño"
        elif gender == 2:
            context.user_data['saludo'] = "La Modelo"
            context.user_data['seno'] = "la señora"
            context.user_data['gen'] = "mujer"
            context.user_data['pana'] = "panameña"
        else:
            update.message.reply_text("Opción no válida. Por favor, elige (1) Hombre o (2) Mujer.")
            return GENDER
    except ValueError:
        update.message.reply_text("Por favor, ingresa un número válido: (1) Hombre, (2) Mujer.")
        return GENDER

    reply_text = f"Perfecto, ahora, por favor ingresa el nombre:"
    update.message.reply_text(reply_text)
    return NAME

# Función para manejar el nombre
def get_name(update, context):
    context.user_data['name'] = update.message.text
    reply_text = f"Gracias, {context.user_data['name']}. Ahora ingresa la cédula con guiones:"
    update.message.reply_text(reply_text)
    return CEDULA

# Función para manejar la cédula
def get_cedula(update, context):
    cedula_input = update.message.text
    cedula_en_letras = convertir_cedula_a_letras(cedula_input)  # Conserva los guiones
    context.user_data['cedula'] = cedula_en_letras
    context.user_data['cedulanumero'] = cedula_input

    reply_text = "Ingresa la fecha en que se realizará la sesión (por ejemplo, sábado 9 de septiembre de 2023):"
    update.message.reply_text(reply_text)
    return DATE

# Función para manejar la fecha
def get_date(update, context):
    context.user_data['date'] = update.message.text
    reply_text = "Ingresa la hora de inicio (por ejemplo, 10:00 a.m.):"
    update.message.reply_text(reply_text)
    return START_TIME

# Función para manejar la hora de inicio
def get_start_time(update, context):
    context.user_data['start_time'] = update.message.text
    reply_text = "Ingresa la hora de finalización (por ejemplo, 12:00 p.m.):"
    update.message.reply_text(reply_text)
    return END_TIME

# Función para manejar la hora de finalización
def get_end_time(update, context):
    context.user_data['end_time'] = update.message.text
    reply_text = "Ingresa el día de la firma del contrato (por ejemplo, martes 5 de septiembre de 2023):"
    update.message.reply_text(reply_text)
    return SIGN_DATE

# Función para manejar la fecha de firma
def get_sign_date(update, context):
    context.user_data['sign_date'] = update.message.text
    reply_text = "¡Gracias! Procesando la información y generando el contrato..."

    try:
        # Diccionario con las variables a reemplazar
        variables_a_reemplazar = {
            "NOMBREM": context.user_data['name'],
            "CEDULAM": context.user_data['cedula'] + "(" + context.user_data['cedulanumero'] + ")",
            "DIAS": context.user_data['date'],
            "HORAI": context.user_data['start_time'],
            "HORAF": context.user_data['end_time'],
            "DIAF": context.user_data['sign_date'],
            "NUMEROCEDULA": context.user_data['cedulanumero'],
            "señora": context.user_data['seno'],
            "mujer": context.user_data['gen'],
            "panameña": context.user_data['pana']
        }

        # Cargamos el documento Word desde la URL
        documento = Document(onedrive_word_plantilla)

        # Recorre el contenido del documento
        for paragraph in documento.paragraphs:
            for variable, valor in variables_a_reemplazar.items():
                if variable in paragraph.text:
                    nuevo_texto = paragraph.text.replace(variable, valor)
                    paragraph.clear()  # Borra el párrafo actual
                    paragraph.add_run(nuevo_texto)  # Agrega el nuevo texto al párrafo

        # Guarda el documento con los cambios
        docx_filename = f"Contrato_para_{context.user_data['name']}.docx"
        documento.save(docx_filename)

        # Convierte el documento Word a PDF
        pdf_filename = docx_filename.replace(".docx", ".pdf")
        docx2pdf.convert(docx_filename, pdf_filename)

        # Guarda el PDF en OneDrive
        guardar_en_onedrive(pdf_filename, onedrive_pdf_url)

        # Elimina el archivo temporal (.docx)
        os.remove(docx_filename)

        # Elimina el archivo PDF temporal
        os.remove(pdf_filename)

        update.message.reply_text("¡Listo! El contrato ha sido generado en PDF y enviado a OneDrive. ¡Que tengas un buen día!")

        return ConversationHandler.END
    except Exception as e:
        update.message.reply_text(
            "Ocurrió un error al generar el contrato. Por favor, intenta nuevamente más tarde o verifica los datos ingresados.")
        print(f"Error en la generación del contrato: {str(e)}")
        return ConversationHandler.END


# Función para mostrar la ayuda
def ayuda(update, context):
    mensaje_ayuda = """
    ¡Hola! Soy tu bot de contrato. Aquí están los comandos disponibles:

    /start - Iniciar la conversación para generar un contrato.
    /ayuda - Mostrar esta ayuda.
    """
    update.message.reply_text(mensaje_ayuda)

# Función principal para iniciar el bot
def main():
    # Crea una instancia de Updater
    updater = Updater(token='TU_TOKEN_AQUI', use_context=True)

    # Crea un manejador de conversación
    conversation_handler = ConversationHandler(
        entry_points=[CommandHandler('start', start)],
        states={
            GENDER: [MessageHandler(Filters.text & ~Filters.command, choose_gender)],
            NAME: [MessageHandler(Filters.text & ~Filters.command, get_name)],
            CEDULA: [MessageHandler(Filters.text & ~Filters.command, get_cedula)],
            DATE: [MessageHandler(Filters.text & ~Filters.command, get_date)],
            START_TIME: [MessageHandler(Filters.text & ~Filters.command, get_start_time)],
            END_TIME: [MessageHandler(Filters.text & ~Filters.command, get_end_time)],
            SIGN_DATE: [MessageHandler(Filters.text & ~Filters.command, get_sign_date)]
        },
        fallbacks=[CommandHandler('ayuda', ayuda)]
    )

    dispatcher = updater.dispatcher
    dispatcher.add_handler(conversation_handler)

    # Agrega un manejador para el comando /ayuda
    dispatcher.add_handler(CommandHandler('ayuda', ayuda))

    # Agrega un manejador para el comando /stop
    dispatcher.add_handler(CommandHandler('stop', stop))

    # Inicia el bot
    updater.start_polling()
    updater.idle()

if __name__ == '__main__':
    print("Bot iniciado.")
    main()